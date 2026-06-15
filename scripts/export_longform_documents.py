#!/usr/bin/env python3
"""Export a longform Markdown report to DOCX and PDF."""

from __future__ import annotations

import argparse
import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BODY_FONT = "Microsoft YaHei"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export longform Markdown to DOCX and PDF.")
    parser.add_argument("markdown")
    parser.add_argument("--docx", required=True)
    parser.add_argument("--pdf", required=True)
    parser.add_argument("--subtitle", default="综合命盘长文报告")
    return parser.parse_args()


def markdown_runs(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def set_east_asia_font(run, font_name: str = BODY_FONT) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)


def configure_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.208


def add_docx_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    set_east_asia_font(run)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(32, 55, 72)

    if subtitle.strip():
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(36)
        run = sub.add_run(subtitle)
        set_east_asia_font(run)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_section(WD_SECTION.NEW_PAGE)


def export_docx(markdown: str, output: Path, subtitle: str) -> None:
    lines = markdown.splitlines()
    title = markdown_runs(lines[0].lstrip("# ").strip()) if lines and lines[0].startswith("# ") else output.stem
    doc = Document()
    configure_docx_styles(doc)
    add_docx_title(doc, title, subtitle)

    for raw in lines[1:]:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_paragraph(markdown_runs(line[4:]), style="Heading 2")
        elif line.startswith("## "):
            doc.add_paragraph(markdown_runs(line[3:]), style="Heading 1")
        elif line.startswith("# "):
            doc.add_paragraph(markdown_runs(line[2:]), style="Heading 1")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(markdown_runs(line[2:]))
            set_east_asia_font(run)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(markdown_runs(line))
            set_east_asia_font(run)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def register_pdf_fonts() -> tuple[str, str]:
    body_path = Path(r"C:\Windows\Fonts\simfang.ttf")
    heading_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    body_font = "XuanxueBody"
    heading_font = "XuanxueHeading"
    pdfmetrics.registerFont(TTFont(body_font, str(body_path)))
    pdfmetrics.registerFont(TTFont(heading_font, str(heading_path)))
    return body_font, heading_font


def para_text(text: str) -> str:
    return html.escape(markdown_runs(text))


def export_pdf(markdown: str, output: Path, subtitle: str) -> None:
    body_font, heading_font = register_pdf_fonts()
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontName=heading_font,
            fontSize=22,
            leading=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#203748"),
            spaceAfter=10,
            wordWrap="CJK",
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            parent=base["Normal"],
            fontName=body_font,
            fontSize=11,
            leading=17,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#555555"),
            spaceAfter=36,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "Heading1",
            parent=base["Heading1"],
            fontName=heading_font,
            fontSize=15,
            leading=22,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=16,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "Heading2",
            parent=base["Heading2"],
            fontName=heading_font,
            fontSize=12.5,
            leading=18,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=10,
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName=body_font,
            fontSize=10.5,
            leading=17,
            alignment=TA_LEFT,
            firstLineIndent=18,
            spaceAfter=7,
            wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            parent=base["BodyText"],
            fontName=body_font,
            fontSize=10.5,
            leading=17,
            leftIndent=18,
            firstLineIndent=-10,
            alignment=TA_LEFT,
            spaceAfter=5,
            wordWrap="CJK",
        ),
    }
    lines = markdown.splitlines()
    title = markdown_runs(lines[0].lstrip("# ").strip()) if lines and lines[0].startswith("# ") else output.stem
    story = [Spacer(1, 96), Paragraph(para_text(title), styles["title"])]
    if subtitle.strip():
        story.append(Paragraph(para_text(subtitle), styles["subtitle"]))
    story.append(PageBreak())
    for raw in lines[1:]:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            story.append(Paragraph(para_text(line[4:]), styles["h2"]))
        elif line.startswith("## "):
            story.append(Paragraph(para_text(line[3:]), styles["h1"]))
        elif line.startswith("# "):
            story.append(Paragraph(para_text(line[2:]), styles["h1"]))
        elif line.startswith("- "):
            story.append(Paragraph("• " + para_text(line[2:]), styles["bullet"]))
        else:
            story.append(Paragraph(para_text(line), styles["body"]))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output),
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
        title=title,
        author="xuanxue-console",
    )
    doc.build(story)


def main() -> None:
    args = parse_args()
    markdown = Path(args.markdown).read_text(encoding="utf-8")
    export_docx(markdown, Path(args.docx), args.subtitle)
    export_pdf(markdown, Path(args.pdf), args.subtitle)


if __name__ == "__main__":
    main()
